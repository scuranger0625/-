from PyPDF2 import PdfReader
from nltk.tokenize import sent_tokenize
import openai
import os
from docx import Document
from docx.shared import Inches
import time

# 開始計時
start_time = time.time()

# 設置 OpenAI API 金鑰
openai.api_key = "sk-proj-tXsJfPWYz4cI3CLAObzIT3BlbkFJ9POG15DPqqkeEFaNpYR3"

def get_escaped_path(raw_path):
    """
    從原始路徑中獲取轉義後的檔案路徑。
    """
    escaped_path = r"{}".format(raw_path)
    return escaped_path

def translate_pdf_to_chinese(pdf_path):
    # 讀取 PDF 檔案
    with open(pdf_path, 'rb') as f:
        reader = PdfReader(f)
        number_of_pages = len(reader.pages)
        chunks = []

        for i in range(number_of_pages):
            page = reader.pages[i]
            text = page.extract_text()
            sentences = sent_tokenize(text)
            input_sentences = ''

            for sentence in sentences:
                input_sentences += sentence
                if len(input_sentences) > 1000:
                    chunks.append(input_sentences)
                    input_sentences = ''
            chunks.append(input_sentences)

        # 使用 OpenAI API 完成翻譯
        translated_text = ""
        for chunk in chunks:
            prompt = "翻譯以下文字為中文：\n\n" + chunk
            response = openai.ChatCompletion.create(
                model="gpt-3.5-turbo",
                messages=[
                    {"role": "system", "content": "翻譯以下文字為中文："},
                    {"role": "user", "content": chunk}
                ],
                temperature=0.7,
                max_tokens=200,
                top_p=1.0,
                frequency_penalty=0.0,
                presence_penalty=0.0,
                stop=["\n"]
            )
            translation = response.choices[0].message['content']
            translated_text += translation + "\n"

        return translated_text

def save_to_word(text, title, pdf_path, output_dir):
    doc = Document()
    doc.add_heading(title, level=1)
    
    # 添加文字內容
    doc.add_paragraph(text)
        
    # 保存Word檔案
    doc.save(os.path.join(output_dir, f"{title}.docx"))

# 文件路徑
raw_path = r"C:\Users\Leon\Desktop\消費者行為期末報告\The Role of Habit in Post-Adoption Switching of Personal Information Technologies An Empirical Investigation.pdf"
pdf_path = get_escaped_path(raw_path)
translated_text = translate_pdf_to_chinese(pdf_path)

# 提取檔案名稱作為標題
title = os.path.basename(raw_path).split('.')[0]

# 保存到桌面
desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
save_to_word(translated_text, title, pdf_path, desktop_path)


# 結束計時
end_time = time.time()
elapsed_time = end_time - start_time
print(f"Word檔案保存完成，共花費 {elapsed_time:.2f} 秒")